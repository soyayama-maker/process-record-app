import io
from datetime import datetime
from pathlib import Path

import streamlit as st
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

st.set_page_config(
    page_title="精神看護論実習 プロセスレコード練習アプリ v8",
    page_icon="📝",
    layout="wide",
)

TEMPLATE_FILE_NAME = "精神看護論実習記録プロセスレコード.docx"

SCENES = {
    "場面1：話しかけても返答が少ない患者": {
        "theme": "沈黙や短い返答に対して、焦らず、患者のペースを尊重しながら関わる。",
        "setting": "精神科病棟での実習3日目。学生は、受け持ち患者であるAさんに日中の様子を聞こうと思い、デイルームで声をかけた。Aさんは椅子に座り、窓の外を見ている。表情は乏しく、声をかけてもあまり視線を合わせない。",
        "turns": [
            {
                "situation": "",
                "patient": "別に……話すことはありません。",
                "options": {
                    "A": {"student": "そうなんですね。無理に話さなくても大丈夫です。少しここにいてもよいですか。", "reaction": "……別に、いいですけど。"},
                    "B": {"student": "でも、少しは話した方が気分転換になると思いますよ。", "reaction": "そういうのが、しんどいです。"},
                    "C": {"student": "では、また後で来たほうがいいですか。", "reaction": "……どちらでもいいです。"},
                    "D": {"student": "話したくない理由は何ですか。", "reaction": "理由って言われても……。別にないです。"},
                },
            },
            {
                "situation": "少し沈黙があった後、Aさんは視線を窓の外に向けたまま、小さな声で話した。",
                "patient": "ここにいても、何をしたらいいのか分からないんです。",
                "options": {
                    "A": {"student": "何をしたらいいのか分からない感じがあるんですね。", "reaction": "……そうですね。毎日、同じような感じで。"},
                    "B": {"student": "作業療法やレクリエーションに参加してみたらどうですか。", "reaction": "そういうの、苦手なんです。"},
                    "C": {"student": "入院中なので、今は治療に集中するしかないですね。", "reaction": "……そうですね。"},
                    "D": {"student": "何もすることがないと、つらくなりますよね。", "reaction": "つらいというか……空っぽな感じです。"},
                },
            },
            {
                "situation": "",
                "patient": "前は、家で音楽を聴いたりしていたんですけど、今はそういう気にもならなくて。",
                "options": {
                    "A": {"student": "音楽を聴くことが、以前はAさんにとって過ごし方の一つだったんですね。", "reaction": "……そうですね。少し落ち着く感じはありました。"},
                    "B": {"student": "また音楽を聴けるようになるといいですね。", "reaction": "そうですね……でも、今はまだ無理です。"},
                    "C": {"student": "今は気分が落ちているから、仕方ないと思います。", "reaction": "……はい。"},
                    "D": {"student": "どんな音楽を聴いていたんですか。", "reaction": "昔の曲です。詳しくはないですけど、静かな曲を聴いていました。"},
                },
            },
        ],
    },
    "場面2：退院への不安を話す患者": {
        "theme": "退院への不安を安易に励ますのではなく、不安の背景を受け止め、希望やストレングスに目を向ける。",
        "setting": "精神科病棟での実習5日目。学生は、受け持ち患者であるBさんとデイルームで話をしている。Bさんは数週間後に退院の予定があるが、退院後の生活について不安を感じている様子がある。Bさんは少しうつむきながら、ぽつりと話し始めた。",
        "turns": [
            {
                "situation": "",
                "patient": "退院しても、どうせまた悪くなると思います。",
                "options": {
                    "A": {"student": "退院後にまた調子が悪くなるのではないかと心配されているんですね。", "reaction": "……はい。前も退院した後に、しんどくなったので。"},
                    "B": {"student": "そんなふうに考えすぎなくても大丈夫だと思いますよ。", "reaction": "そう言われても……不安なものは不安です。"},
                    "C": {"student": "退院が決まっているということは、よくなっているということではないですか。", "reaction": "病院ではそうかもしれないですけど、家に帰ったら分かりません。"},
                    "D": {"student": "また悪くなると思うのは、どんな時ですか。", "reaction": "夜に一人でいる時とか……いろいろ考えてしまいます。"},
                },
            },
            {
                "situation": "Bさんは少し視線を落とし、手元を見ながら話を続けた。",
                "patient": "家に帰っても、ちゃんと生活できる自信がないんです。",
                "options": {
                    "A": {"student": "生活できる自信がないと感じているんですね。", "reaction": "はい……また生活が崩れるんじゃないかと思って。"},
                    "B": {"student": "退院後は訪問看護などもありますし、きっと大丈夫だと思います。", "reaction": "支援があるのは分かるんですけど、それでも不安です。"},
                    "C": {"student": "具体的には、どのようなことが一番心配ですか。", "reaction": "薬を飲み忘れたり、寝られなくなったりすることです。"},
                    "D": {"student": "入院中にできるようになったこともあるのではないですか。", "reaction": "できるようになったこと……ですか。前よりは、薬は自分で確認するようになりました。"},
                },
            },
            {
                "situation": "",
                "patient": "でも、また家族に迷惑をかけるかもしれないと思うと、退院するのが怖いです。",
                "options": {
                    "A": {"student": "ご家族に迷惑をかけるかもしれないと思うと、退院が怖く感じるんですね。", "reaction": "はい……前も心配をかけたので、また同じことになるのが嫌なんです。"},
                    "B": {"student": "家族はきっと支えてくれると思いますよ。", "reaction": "そうだといいんですけど……またがっかりされるかもしれません。"},
                    "C": {"student": "迷惑をかけないようにするには、どうしたらよいと思いますか。", "reaction": "それが分からないから、不安なんです。"},
                    "D": {"student": "Bさんは、ご家族にどのように過ごしてほしいと思っていますか。", "reaction": "普通に過ごしたいです。できれば、あまり心配されずに生活したいです。"},
                },
            },
        ],
    },
    "場面3：学生に否定的な言葉を向ける患者": {
        "theme": "否定的に聞こえる言葉を向けられたとき、自分の感情に気づきながら、患者の言葉の背景を考えて関わる。",
        "setting": "精神科病棟での実習4日目。学生は、受け持ち患者であるCさんに、午前中の作業療法の様子について話を聞こうと思い、病棟の廊下で声をかけた。Cさんは少し疲れた表情で椅子に座っている。学生が近づくと、Cさんは視線を合わせずに話した。",
        "turns": [
            {
                "situation": "",
                "patient": "学生さんに話しても、どうせ分からないでしょ。",
                "options": {
                    "A": {"student": "そう感じておられるんですね。分からないこともあると思いますが、Cさんの感じていることを教えていただきたいです。", "reaction": "……教えるってほどのことでもないですけど。"},
                    "B": {"student": "そんなことはありません。私も勉強しているので分かります。", "reaction": "勉強して分かるなら、苦労しないですよ。"},
                    "C": {"student": "分からないかもしれませんが、実習なので話を聞かせてください。", "reaction": "実習のためなんですね。"},
                    "D": {"student": "分かってもらえないと感じることが、これまでにもあったのでしょうか。", "reaction": "まあ……何を言っても、結局は病気のせいにされることが多いので。"},
                },
            },
            {
                "situation": "Cさんは少し間を置いて、視線を床に向けたまま話を続けた。",
                "patient": "前にも、話したことをちゃんと聞いてもらえなかったことがあるんです。",
                "options": {
                    "A": {"student": "話したことをちゃんと聞いてもらえなかったと感じたことがあったんですね。", "reaction": "はい……話しても、結局分かってもらえない感じがしました。"},
                    "B": {"student": "でも、看護師さんたちはCさんのためを思って関わっていると思います。", "reaction": "そうかもしれないですけど、そう思えない時もあります。"},
                    "C": {"student": "それはいつのことですか。誰に言われたのですか。", "reaction": "……別に、もういいです。"},
                    "D": {"student": "そういう経験があると、私にも話しにくく感じるかもしれませんね。", "reaction": "そうですね。どうせ同じかなって思ってしまいます。"},
                },
            },
            {
                "situation": "",
                "patient": "本当は、作業療法も行った方がいいんでしょうけど、人が多いところに行くと疲れるんです。",
                "options": {
                    "A": {"student": "人が多い場所に行くと疲れるので、作業療法に行くことにも負担を感じているんですね。", "reaction": "はい……嫌というより、疲れるんです。"},
                    "B": {"student": "でも、参加した方が退院には近づくと思いますよ。", "reaction": "そう言われると、行けない自分が悪いみたいですね。"},
                    "C": {"student": "無理に行かなくてもいいと思います。", "reaction": "そうですか……でも、ずっと行かないのもよくない気がします。"},
                    "D": {"student": "人が多い場所で疲れる時、少し楽に過ごせる方法はありますか。", "reaction": "端の席に座れると、少しはましです。途中で休めるなら行けるかもしれません。"},
                },
            },
        ],
    },
}


def set_run_font(run, font_name="Yu Gothic", size=10.5, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass


def clear_cell(cell):
    cell.text = ""
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""


def set_cell_text(cell, text, bold_first_line=False, font_size=10):
    """Replace cell content with text. Newlines become separate paragraphs."""
    clear_cell(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    lines = str(text or "").split("\n")
    if not lines:
        lines = [""]
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        if idx == 0 and not line:
            continue
        run = p.add_run(line)
        set_run_font(run, size=font_size, bold=(bold_first_line and idx == 0))


def shade_cell(cell, fill="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.first_child_found_in("w:tcW")
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")


def set_paragraph_text(paragraph, text, bold=False, size=10.5):
    paragraph.text = ""
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def circled_number(n: int) -> str:
    circled = {
        1: "①", 2: "②", 3: "③", 4: "④", 5: "⑤",
        6: "⑥", 7: "⑦", 8: "⑧", 9: "⑨", 10: "⑩",
    }
    return circled.get(n, str(n))


def build_by_turn_text(turn_records, key, heading_prefix=None):
    blocks = []
    for rec in turn_records:
        content = (rec.get(key) or "").strip()
        label = rec.get("sequence_label", f"会話{rec.get('turn', '')}")
        blocks.append(f"{label}\n{content}")
    return "\n\n".join(blocks)


def make_fallback_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic"
    style.font.size = Pt(10.5)
    try:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    except Exception:
        pass

    title = doc.add_paragraph()
    set_paragraph_text(title, "精神看護論実習記録⑥（プロセスレコード）", bold=True, size=12)
    info = doc.add_paragraph()
    set_paragraph_text(info, "", size=10.5)

    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.autofit = False
    except Exception:
        pass

    table.cell(0, 1).merge(table.cell(0, 3))
    table.cell(3, 0).merge(table.cell(3, 3))

    # Column widths are approximate. Word may adjust them depending on content.
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, [4.4, 4.4, 4.4, 4.4][idx if idx < 4 else 0])

    return doc


def build_docx_bytes(
    practice_date,
    practice_day,
    student_id,
    student_name,
    scene_record,
    reason_record,
    turn_records,
    learning_record,
    template_path=None,
):
    if template_path and Path(template_path).exists():
        doc = Document(str(template_path))
        # Preserve the provided template as much as possible.
        if len(doc.paragraphs) >= 1:
            set_paragraph_text(doc.paragraphs[0], "精神看護論実習記録⑥（プロセスレコード）", bold=True, size=12)
        if len(doc.paragraphs) >= 2:
            set_paragraph_text(
                doc.paragraphs[1],
                f"記入日　{practice_date or '　　　　'}（実習　{practice_day or '　'}日目）　　　学籍番号　{student_id or '　　　　　　'}　　　学生氏名　{student_name or '　　　　　　　　　'}",
                size=10.5,
            )
        table = doc.tables[0] if doc.tables else None
        if table is None:
            doc = make_fallback_document()
            table = doc.tables[0]
    else:
        doc = make_fallback_document()
        set_paragraph_text(
            doc.paragraphs[1],
            f"記入日　{practice_date or '　　　　'}（実習　{practice_day or '　'}日目）　　　学籍番号　{student_id or '　　　　　　'}　　　学生氏名　{student_name or '　　　　　　　　　'}",
            size=10.5,
        )
        table = doc.tables[0]

    # Fill cells. The designated form has 4 rows: row0 = 1/2, row1 = headings 3-6, row2 = content, row3 = 7.
    # If the template has merged cells, python-docx returns repeated references; writing to the first merged cell is sufficient.
    set_cell_text(table.cell(0, 0), f"1．場面\n{scene_record}", bold_first_line=True, font_size=10)
    set_cell_text(table.cell(0, 1), f"2．この場面を選択した理由\n{reason_record}", bold_first_line=True, font_size=10)

    set_cell_text(table.cell(1, 0), "3．対象者の言動/様子", bold_first_line=True, font_size=10)
    set_cell_text(table.cell(1, 1), "4．学生が感じたこと/考えたこと", bold_first_line=True, font_size=10)
    set_cell_text(table.cell(1, 2), "5．学生の言動", bold_first_line=True, font_size=10)
    set_cell_text(table.cell(1, 3), "6．分析・考察", bold_first_line=True, font_size=10)

    # Light blue header row like the original template often uses after export.
    for c in range(4):
        try:
            shade_cell(table.cell(1, c), "D9EAF7")
        except Exception:
            pass

    patient_text = build_by_turn_text(turn_records, "patient_behavior", "対象者")
    feeling_text = build_by_turn_text(turn_records, "student_feeling", "感情")
    action_text = build_by_turn_text(turn_records, "student_action", "学生")
    analysis_text = build_by_turn_text(turn_records, "analysis", "分析")

    set_cell_text(table.cell(2, 0), patient_text, font_size=9.5)
    set_cell_text(table.cell(2, 1), feeling_text, font_size=9.5)
    set_cell_text(table.cell(2, 2), action_text, font_size=9.5)
    set_cell_text(table.cell(2, 3), analysis_text, font_size=9.5)

    set_cell_text(table.cell(3, 0), f"7．この場面から学んだこと\n{learning_record}", bold_first_line=True, font_size=10)

    # Apply font settings across the document for better Japanese compatibility.
    for p in doc.paragraphs:
        for r in p.runs:
            set_run_font(r, size=r.font.size.pt if r.font.size else 10.5, bold=r.bold)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                for p in cell.paragraphs:
                    for r in p.runs:
                        # Do not override existing bold decisions, only font family.
                        if r.font.size is None:
                            set_run_font(r, size=9.5, bold=r.bold)
                        else:
                            set_run_font(r, size=r.font.size.pt, bold=r.bold)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


st.title("精神看護論実習 プロセスレコード練習アプリ")
st.caption("v8 修正版：対象者の言動→学生が感じたこと／考えたこと→学生の言動の流れで通し番号を表示します。")

st.info(
    "このアプリは、精神看護論実習前に、模擬場面を用いてプロセスレコードの記述を練習するための教材です。"
    "実際の患者情報や個人情報は入力しないでください。"
)

with st.expander("使い方", expanded=False):
    st.write(
        "1. 基本情報を入力します。\n"
        "2. 模擬場面を選択します。\n"
        "3. 患者の発言を読み、学生の返答を選択します。\n"
        "4. 会話ログを確認します。\n"
        "5. 会話ごとにプロセスレコードを記入します。\n"
        "6. 対象者の言動→学生が感じたこと／考えたこと→学生の言動の通し番号を確認し、最後にWordファイルをダウンロードします。"
    )

app_dir = Path(__file__).parent
local_template = app_dir / TEMPLATE_FILE_NAME
template_path = str(local_template) if local_template.exists() else None
if template_path:
    st.success(f"指定様式テンプレートを読み込みます：{TEMPLATE_FILE_NAME}")
else:
    st.warning(
        f"同じフォルダに {TEMPLATE_FILE_NAME} が見つかりません。"
        "この場合も同じ項目構成のWordファイルを作成できますが、正式な指定様式を使う場合は、テンプレートを同じフォルダに置いてから起動してください。"
    )

st.header("基本情報")
col1, col2, col3, col4 = st.columns(4)
with col1:
    practice_date = st.text_input("記入日", placeholder="例：2026年10月15日")
with col2:
    practice_day = st.text_input("実習日目", placeholder="例：3")
with col3:
    student_id = st.text_input("学籍番号")
with col4:
    student_name = st.text_input("氏名")

st.header("場面を選択してください")
scene_name = st.selectbox("模擬場面", list(SCENES.keys()))
scene = SCENES[scene_name]

# 場面を変更した場合は、前の場面の選択状態と記録をリセットする
if st.session_state.get("active_scene") != scene_name:
    st.session_state["active_scene"] = scene_name
    for key in list(st.session_state.keys()):
        if key.startswith("choice_") or key.startswith("prev_choice_"):
            del st.session_state[key]

st.subheader("学習テーマ")
st.write(scene["theme"])
st.subheader("場面設定")
st.write(scene["setting"])

st.header("会話シミュレーション")
st.write("学生の返答を選択すると、患者の反応が表示され、次の会話に進みます。")

if st.button("会話をはじめからやり直す"):
    for key in list(st.session_state.keys()):
        if key.startswith("choice_") or key.startswith("prev_choice_"):
            del st.session_state[key]
    st.rerun()

selected_logs = []
all_turns_completed = True

for i, turn in enumerate(scene["turns"], start=1):
    st.markdown(f"### 会話{i}")
    if turn.get("situation"):
        st.write(f"**患者の様子：** {turn['situation']}")
    st.write(f"**患者：**「{turn['patient']}」")

    option_labels = []
    for key, val in turn["options"].items():
        option_labels.append(f"{key}．{val['student']}")

    choice_key = f"choice_{i}"
    selected_label = st.radio(
        f"学生の返答を選択してください（会話{i}）",
        option_labels,
        index=None,
        key=choice_key,
    )

    if selected_label is None:
        st.info("返答を選択すると、患者の反応が表示されます。")
        all_turns_completed = False
        break

    prev_key = f"prev_choice_{i}"
    if st.session_state.get(prev_key) is not None and st.session_state[prev_key] != selected_label:
        # 前の会話の選択を変更した場合は、その後の会話選択をリセットする
        for j in range(i + 1, len(scene["turns"]) + 1):
            st.session_state.pop(f"choice_{j}", None)
            st.session_state.pop(f"prev_choice_{j}", None)
        st.session_state[prev_key] = selected_label
        st.rerun()
    st.session_state[prev_key] = selected_label

    selected_key = selected_label.split("．", 1)[0]
    selected = turn["options"][selected_key]
    st.success(f"患者の反応：「{selected['reaction']}」")
    selected_logs.append(
        {
            "turn": i,
            "situation": turn.get("situation", ""),
            "patient": turn["patient"],
            "choice_key": selected_key,
            "student": selected["student"],
            "reaction": selected["reaction"],
            "patient_no": circled_number((i - 1) * 3 + 1),
            "feeling_no": circled_number((i - 1) * 3 + 2),
            "student_no": circled_number((i - 1) * 3 + 3),
        }
    )
    st.divider()

if not all_turns_completed:
    st.warning("会話がすべて終わると、会話ログとプロセスレコードの入力欄が表示されます。")
    st.stop()

st.header("あなたが体験した会話")
st.write("記録の順序が分かるように、対象者の言動→学生が感じたこと／考えたこと→学生の言動に通し番号を付けています。")
for log in selected_logs:
    st.markdown(f"**会話{log['turn']}**")
    if log["situation"]:
        st.write(f"患者の様子：{log['situation']}")
    st.write(f"{log['patient_no']} 対象者の言動／様子：{log['patient']}")
    st.write(f"{log['feeling_no']} 学生が感じたこと／考えたこと：この欄はプロセスレコード記入時に入力します。")
    st.write(f"{log['student_no']} 学生の言動：{log['student']}")
    st.caption(f"参考：選択後の患者の反応「{log['reaction']}」")

st.header("プロセスレコードを記入してください")
st.write("各項目の説明は、入力欄の中ではなく項目名の横に表示しています。")
st.info("対象者の言動／様子、学生が感じたこと／考えたこと、学生の言動に①②③…の通し番号を付けています。Word出力でも同じ通し番号が反映されます。")

scene_record = st.text_area(
    "1．場面（どこで、誰と、どのような関わりをした場面かを具体的に書いてください）",
    height=120,
    key="scene_record",
)
reason_record = st.text_area(
    "2．この場面を選択した理由（なぜこの場面が気になったのか、なぜ振り返りたいと思ったのかを書いてください）",
    height=120,
    key="reason_record",
)

turn_records = []
for log in selected_logs:
    i = log["turn"]
    choice_key = log["choice_key"]
    sequence_label = f"会話{i}（{log['patient_no']}〜{log['student_no']}）"
    st.markdown(f"### 会話{i}の記録（{log['patient_no']}対象者の言動／様子 → {log['feeling_no']}学生が感じたこと／考えたこと → {log['student_no']}学生の言動）")
    with st.container(border=True):
        if log["situation"]:
            st.write(f"**患者の様子：** {log['situation']}")
        st.write(f"**{log['patient_no']} 対象者の言動／様子：** {log['patient']}")
        st.write(f"**{log['feeling_no']} 学生が感じたこと／考えたこと：** {log['patient_no']}を受けて感じたこと・考えたことを記入します。")
        st.write(f"**{log['student_no']} 学生の言動：** {log['student']}")
        st.caption(f"参考：学生の言動の後の患者の反応「{log['reaction']}」。分析・考察を書く際の参考にしてください。")

        patient_behavior = st.text_area(
            f"3-{i}．対象者の言動／様子（{log['patient_no']}対象者の発言・様子のみを具体的に書いてください）",
            value=(
                (f"患者の様子：{log['situation']}\n" if log["situation"] else "")
                + f"{log['patient_no']} 対象者の言動／様子：{log['patient']}"
            ),
            height=120,
            key=f"patient_behavior_{scene_name}_{i}_{choice_key}",
        )
        student_feeling = st.text_area(
            f"4-{i}．学生が感じたこと／考えたこと（{log['feeling_no']}{log['patient_no']}を受けて、自分が感じたこと・考えたこと・戸惑ったことを書いてください）",
            value=f"{log['feeling_no']} ",
            height=120,
            key=f"student_feeling_{scene_name}_{i}_{choice_key}",
        )
        student_action = st.text_area(
            f"5-{i}．学生の言動（{log['student_no']}自分が実際にどのように声をかけたか、どのような態度で関わったかを書いてください）",
            value=f"{log['student_no']} 学生の言動：{log['student']}",
            height=110,
            key=f"student_action_{scene_name}_{i}_{choice_key}",
        )
        analysis_record = st.text_area(
            f"6-{i}．分析・考察（{log['patient_no']}〜{log['student_no']}の流れから、対象者の言動の背景や自分の関わりの影響を考えてください）",
            value=f"{log['patient_no']}〜{log['student_no']} ",
            height=150,
            key=f"analysis_record_{scene_name}_{i}_{choice_key}",
        )
        turn_records.append(
            {
                "turn": i,
                "sequence_label": sequence_label,
                "patient_behavior": patient_behavior,
                "student_feeling": student_feeling,
                "student_action": student_action,
                "analysis": analysis_record,
            }
        )

learning_record = st.text_area(
    "7．この場面から学んだこと（この場面を通して学んだこと、次の関わりで意識したいことを書いてください）",
    height=130,
    key="learning_record",
)

st.header("自己チェック")
check_items = [
    "対象者の言葉や様子を具体的に書けている",
    "自分の感情や考えを書けている",
    "自分の言動を具体的に書けている",
    "対象者の言動の背景を考えている",
    "自分の関わりが対象者に与えた影響を考えている",
    "次の関わりに活かせる学びを書けている",
    "「うまくいった／うまくいかなかった」だけで終わらず、理由を考えている",
]
for item in check_items:
    st.checkbox(item, key=f"check_{item}")

st.header("Wordファイルのダウンロード")
st.write("入力内容を、精神看護論実習記録⑥（プロセスレコード）のWord様式に落とし込んでダウンロードします。")

word_bytes = build_docx_bytes(
    practice_date=practice_date,
    practice_day=practice_day,
    student_id=student_id,
    student_name=student_name,
    scene_record=scene_record,
    reason_record=reason_record,
    turn_records=turn_records,
    learning_record=learning_record,
    template_path=template_path,
)

safe_student_id = student_id if student_id else "student"
file_name = f"精神看護論実習記録⑥_プロセスレコード_{safe_student_id}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
st.download_button(
    label="Wordファイルでダウンロード",
    data=word_bytes,
    file_name=file_name,
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
)

st.warning("このアプリは模擬事例による練習用です。実際の患者情報や個人情報は入力しないでください。")
